"""Report generation for scheduled and on-demand automation jobs."""

from __future__ import annotations

import asyncio
import re
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from assistant.automation.models import ReportFormat, ReportJob, ReportResult, ReportSource
from assistant.utils.logging import get_logger


LOGGER = get_logger("report_generator")
SUPPORTED_FOLDER_SUFFIXES = {".txt", ".md", ".py", ".docx", ".pdf"}


class ReportGenerator:
    def __init__(self, config, model_provider, tool_registry) -> None:
        self.config = config
        self.model_provider = model_provider
        self.tool_registry = tool_registry
        self.memory_manager = None
        self.delivery = None

    def bind_runtime(self, *, memory_manager=None, delivery=None) -> None:
        self.memory_manager = memory_manager
        self.delivery = delivery

    async def generate(self, job: ReportJob) -> ReportResult:
        sources = await self._collect_sources(job)
        report_content = await self._summarize(job.topic, sources)
        rendered = await self._format_output(report_content, job.output_format.value)
        save_path = self._resolve_save_path(job)
        await self._write_output(save_path, rendered)
        byte_size = save_path.stat().st_size if save_path.exists() else 0
        return ReportResult(
            job_id=job.job_id,
            topic=job.topic,
            save_path=str(save_path.resolve()),
            format=job.output_format.value,
            byte_size=byte_size,
            generated_at=datetime.now(timezone.utc).isoformat(),
            summary_preview=report_content[:500],
        )

    async def _collect_sources(self, job: ReportJob) -> list[str]:
        sources: list[str] = []
        source_type = job.source_type
        if source_type in {ReportSource.folder, ReportSource.mixed} and job.source_path:
            sources.extend(await self._collect_folder_sources(job.source_path))
        if source_type in {ReportSource.web_search, ReportSource.mixed}:
            sources.extend(await self._collect_web_sources(job.topic))
        if source_type in {ReportSource.memory, ReportSource.mixed}:
            sources.extend(await self._collect_memory_sources())
        return sources or [f"No sources were available for topic '{job.topic}'."]

    async def _summarize(self, topic: str, sources: list[str]) -> str:
        joined_sources = "\n\n".join(sources)
        prompt = (
            f"You are a report writer. Topic: {topic}\n"
            "Sources below. Write a structured report with:\n"
            "- Executive Summary (2-3 sentences)\n"
            "- Key Points (bullet list)\n"
            "- Details (prose paragraphs)\n"
            "- Conclusion\n"
            "Use Markdown formatting.\n"
            "Sources:\n"
            f"{joined_sources}"
        )
        chunks: list[str] = []
        async for response in self.model_provider.complete(
            messages=[{"role": "user", "content": prompt}],
            system="You write concise, well-structured reports from provided source material.",
            tools=[],
            stream=False,
        ):
            if response.text:
                chunks.append(response.text)
        content = "".join(chunks).strip()
        return content or f"# {topic}\n\nNo report content could be generated."

    async def _format_output(self, content: str, fmt: str) -> bytes | str:
        if fmt == ReportFormat.markdown.value:
            return content
        if fmt == ReportFormat.txt.value:
            return self._strip_markdown(content)
        if fmt == ReportFormat.docx.value:
            try:
                from docx import Document  # type: ignore
            except Exception:
                return content
            document = Document()
            for line in content.splitlines():
                stripped = line.strip()
                if not stripped:
                    document.add_paragraph("")
                    continue
                if stripped.startswith("# "):
                    document.add_heading(stripped[2:].strip(), level=1)
                elif stripped.startswith("## "):
                    document.add_heading(stripped[3:].strip(), level=2)
                elif stripped.startswith("- "):
                    document.add_paragraph(stripped[2:].strip(), style="List Bullet")
                else:
                    document.add_paragraph(self._strip_markdown(stripped))
            buffer = BytesIO()
            document.save(buffer)
            return buffer.getvalue()
        if fmt == ReportFormat.pdf.value:
            pdf_bytes = await self._build_pdf_bytes(content)
            if pdf_bytes is not None:
                return pdf_bytes
            return content
        return content

    def _resolve_save_path(self, job: ReportJob) -> Path:
        reports_root = self.config.agent.workspace_dir / self.config.reports.reports_subdir
        reports_root.mkdir(parents=True, exist_ok=True)
        extension = self._extension_for_format(job.output_format.value)
        if not job.save_path:
            slug = self._slugify(job.topic)
            candidate = reports_root / f"{datetime.now(timezone.utc).strftime('%Y%m%d')}_{slug}.{extension}"
        else:
            raw_path = Path(job.save_path)
            if raw_path.is_absolute():
                candidate = raw_path
            else:
                candidate = reports_root / raw_path
            if candidate.suffix.lower() != f".{extension}":
                candidate = candidate.with_suffix(f".{extension}")
        candidate.parent.mkdir(parents=True, exist_ok=True)
        resolved = candidate
        counter = 1
        while resolved.exists():
            resolved = candidate.with_name(f"{candidate.stem}_{counter}{candidate.suffix}")
            counter += 1
        return resolved.resolve()

    async def _deliver(self, result: ReportResult, job: ReportJob) -> None:
        if self.delivery is None:
            return
        preview = result.summary_preview[:3000]
        path = Path(result.save_path)
        deliver_via = str(job.deliver_via or self.config.reports.default_deliver_via).lower()
        if deliver_via in {"telegram", "all"}:
            await self.delivery.send_text(
                job.user_id,
                f"Report ready: {job.topic}\n\n{preview}",
                channel_name="telegram",
            )
            if path.exists() and path.stat().st_size < 10 * 1024 * 1024:
                await self.delivery.send_file(job.user_id, path, channel_name="telegram", caption=f"Report: {job.topic}")
        if deliver_via in {"memory", "all"} and self.memory_manager is not None:
            await self.memory_manager.write_long_term(
                f"Report: {job.topic}",
                f"Saved to {result.save_path}\n\n{result.summary_preview}",
            )
        if deliver_via in {"all", "telegram", "memory"}:
            return
        # "file" intentionally does not notify.

    async def _collect_folder_sources(self, source_path: str) -> list[str]:
        root = Path(source_path).expanduser()
        max_files = int(getattr(self.config.reports, "max_source_files", 20))
        max_chars = int(getattr(self.config.reports, "max_source_chars_per_file", 3000))
        collected: list[str] = []
        for path in self._iter_supported_files(root):
            if len(collected) >= max_files:
                break
            text = await self._read_source_file(path)
            if not text:
                continue
            collected.append(f"## {path}\n{text[:max_chars]}")
        return collected

    async def _collect_web_sources(self, topic: str) -> list[str]:
        if self.tool_registry is None or not self.tool_registry.has("web_search"):
            return []
        try:
            result = await self.tool_registry.dispatch("web_search", {"query": topic, "limit": 5})
        except Exception as exc:
            LOGGER.warning("report_web_search_failed", topic=topic, error=str(exc))
            return []
        snippets = result.get("results", [])
        return [f"## Web Result {index + 1}\n{item}" for index, item in enumerate(snippets[:5])]

    async def _collect_memory_sources(self) -> list[str]:
        if self.memory_manager is None:
            return []
        long_term = await self.memory_manager.read_long_term()
        daily = await self.memory_manager.read_today_and_yesterday()
        parts = []
        if long_term.strip():
            parts.append(f"## Long-Term Memory\n{long_term[:6000]}")
        if daily.strip():
            parts.append(f"## Recent Daily Memory\n{daily[:6000]}")
        return parts

    async def _read_source_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                return await self._read_pdf(path)
            if suffix == ".docx":
                return await self._read_docx(path)
            return await asyncio.to_thread(path.read_text, encoding="utf-8", errors="ignore")
        except Exception as exc:
            LOGGER.warning("report_source_read_failed", path=str(path), error=str(exc))
            return ""

    async def _read_pdf(self, path: Path) -> str:
        if self.tool_registry is not None and self.tool_registry.has("pdf_extract"):
            try:
                result = await self.tool_registry.dispatch("pdf_extract", {"path": str(path)})
                return str(result.get("content", ""))
            except Exception as exc:
                LOGGER.warning("report_pdf_extract_failed", path=str(path), error=str(exc))
        try:
            import pdfplumber  # type: ignore
        except Exception:
            return ""

        def _extract() -> str:
            parts: list[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages[:10]:
                    parts.append(page.extract_text() or "")
            return "\n\n".join(part for part in parts if part.strip())

        return await asyncio.to_thread(_extract)

    async def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore
        except Exception:
            return await asyncio.to_thread(path.read_text, encoding="utf-8", errors="ignore")

        def _extract() -> str:
            document = Document(path)
            return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

        return await asyncio.to_thread(_extract)

    async def _build_pdf_bytes(self, content: str) -> bytes | None:
        try:
            from reportlab.lib.pagesizes import letter  # type: ignore
            from reportlab.pdfgen import canvas  # type: ignore
        except Exception:
            try:
                from fpdf import FPDF  # type: ignore
            except Exception:
                return None
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in self._strip_markdown(content).splitlines():
                pdf.multi_cell(0, 8, line or " ")
            return pdf.output(dest="S").encode("latin-1", errors="ignore")

        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        y = height - 50
        for line in self._strip_markdown(content).splitlines():
            if y < 50:
                pdf.showPage()
                y = height - 50
            pdf.drawString(50, y, line[:110] or " ")
            y -= 15
        pdf.save()
        return buffer.getvalue()

    async def _write_output(self, path: Path, rendered: bytes | str) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        if isinstance(rendered, bytes):
            await asyncio.to_thread(path.write_bytes, rendered)
        else:
            await asyncio.to_thread(path.write_text, rendered, encoding="utf-8")

    def _iter_supported_files(self, root: Path) -> list[Path]:
        if not root.exists():
            return []
        files = [path for path in root.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_FOLDER_SUFFIXES]
        files.sort(key=lambda item: str(item).lower())
        return files

    def _extension_for_format(self, fmt: str) -> str:
        return {
            ReportFormat.markdown.value: "md",
            ReportFormat.txt.value: "txt",
            ReportFormat.docx.value: "docx",
            ReportFormat.pdf.value: "pdf",
        }.get(fmt, "md")

    def _slugify(self, topic: str) -> str:
        slug = re.sub(r"[^a-z0-9_]+", "", re.sub(r"\s+", "_", topic.lower()).strip("_"))
        return slug or "report"

    def _strip_markdown(self, content: str) -> str:
        plain = re.sub(r"^\s{0,3}#{1,6}\s*", "", content, flags=re.MULTILINE)
        plain = re.sub(r"[*_`>#-]+", "", plain)
        plain = re.sub(r"\n{3,}", "\n\n", plain)
        return plain.strip()
